import datetime

from docx import Document 
from docx.shared import Inches

class Papertime:
    def __init__(self, straps: list, strap_colour = None):
        self.straps = straps 
        self.strap_colour = strap_colour 
        self.date = datetime.date.today() 
        self.strap_colour = strap_colour

    def make_document(self) -> None:
        document = Document() 
        title_colour = self.strap_colour if self.strap_colour else 'All'
        document.add_heading(f'{self.date} -- {title_colour}') 

        # TODO: fix the argument placement within the map functions
        # DEVIDE BY URGENCY
        # -----------------
        urgent_straps = list(filter(lambda strap: strap.urgent_bool, self.straps)) 
        non_urgent_straps = list(filter(lambda strap: not strap.urgent_bool, self.straps)) 
        # -----------------
        # DEVIDE BY COLOUR 
        # -----------------
        urgent_red_straps = list(filter(lambda strap: strap.colour.lower() == 'red', urgent_straps))
        urgent_blue_straps = list(filter(lambda strap: strap.colour.lower() == 'navy', urgent_straps)) 
        # print('urgent_red_straps: {}'.format([strap for strap in urgent_red_straps])) 
        # print('urgent_blue_straps: {}'.format([strap.colour.lower() for strap in urgent_blue_straps]))    

        non_urgent_red_straps = list(filter(lambda strap: strap.colour.lower() == 'red', non_urgent_straps))
        non_urgent_blue_straps = list(filter(lambda strap: strap.colour.lower() == 'navy', non_urgent_straps)) 
        # -----------------
        # DEVIDE BY BUCKLE TYPE 
        # ----------------- 
        urgent_red_straps_NORMLE_BUCKLE = list(filter(lambda strap: strap.buckle_type.lower() == 'new', urgent_red_straps))
        urgent_red_straps_OLD_BUCKLE = list(filter(lambda strap: strap.buckle_type.lower() == 'old', urgent_red_straps))
        # print('strap in urgent_red_straps_NORMLE_BUCKLE: {}'.format([strap for strap in urgent_red_straps_NORMLE_BUCKLE]))
        # print('strap in urgent_red_straps_OLD_BUCKLE: {}'.format([strap for strap in urgent_red_straps_OLD_BUCKLE]))

        urgent_blue_straps_NORMLE_BUCKLE = list(filter(lambda strap: strap.buckle_type.lower() == 'new', urgent_blue_straps))
        urgent_blue_straps_OLD_BUCKLE = list(filter(lambda strap: strap.buckle_type.lower() == 'old', urgent_blue_straps))
        # print('strap in urgent_blue_straps_NORMLE_BUCKLE: {}'.format([strap for strap in urgent_blue_straps_NORMLE_BUCKLE]))
        # print('strap in urgent_blue_straps_OLD_BUCKLE: {}'.format([strap for strap in urgent_blue_straps_OLD_BUCKLE]))

        non_urgent_red_straps_NORMLE_BUCKLE = list(filter(lambda strap: strap.buckle_type.lower() == 'new', non_urgent_red_straps))
        non_urgent_red_straps_OLD_BUCKLE = list(filter(lambda strap: strap.buckle_type.lower() == 'old', non_urgent_red_straps))
        # print('strap in non_urgent_red_straps_NORMLE_BUCKLE: {}'.format([strap for strap in non_urgent_red_straps_NORMLE_BUCKLE]))
        # print('strap in non_urgent_red_straps_OLD_BUCKLE: {}'.format([strap for strap in non_urgent_red_straps_OLD_BUCKLE]))

        non_urgent_blue_straps_NORMLE_BUCKLE = list(filter(lambda strap: strap.buckle_type.lower() == 'new', non_urgent_blue_straps))
        non_urgent_blue_straps_OLD_BUCKLE = list(filter(lambda strap: strap.buckle_type.lower() == 'old', non_urgent_blue_straps))
        # print('strap in non_urgent_blue_straps_NORMLE_BUCKLE: {}'.format([strap for strap in non_urgent_blue_straps_NORMLE_BUCKLE]))
        # print('strap in non_urgent_blue_straps_OLD_BUCKLE: {}'.format([strap for strap in non_urgent_blue_straps_OLD_BUCKLE]))
        # -----------------

        blue_index = 1
        red_index = 1
        
        there_are_urgent_straps = len([l for l in self.straps if l.urgent_bool]) 
        there_are_non_urgent_straps = len([l for l in self.straps if not l.urgent_bool])

        if there_are_urgent_straps:
            document.add_heading('URGENT')
            p = document.add_paragraph() 

        for strap in urgent_blue_straps_NORMLE_BUCKLE:
            p.add_run(f"{blue_index}B | {strap.qty}X  '{strap.name}' \n").bold = True
            blue_index += 1

        for strap in urgent_blue_straps_OLD_BUCKLE:
            p.add_run(f"{blue_index}B | {strap.qty}X '{strap.name}' {strap.buckle_type} Buckle \n").bold = True
            blue_index += 1 
            
        for strap in urgent_red_straps_NORMLE_BUCKLE:
            p.add_run(f"{red_index}R | {strap.qty}X '{strap.name}' \n") 
            red_index += 1 
 
        for strap in urgent_red_straps_OLD_BUCKLE:
            p.add_run(f"{red_index}R | {strap.qty}X '{strap.name}' {strap.buckle_type} Buckle \n")
            red_index += 1

        if there_are_non_urgent_straps:
            print('yay 2')
            document.add_heading('NORMAL') 
            p2 = document.add_paragraph()

        for strap in non_urgent_blue_straps_NORMLE_BUCKLE:
            p2.add_run(f"{blue_index}B | {strap.qty}X  '{strap.name}' \n").bold = True
            blue_index += 1

        for strap in non_urgent_blue_straps_OLD_BUCKLE:
            p2.add_run(f"{blue_index}B | {strap.qty}X '{strap.name}' {strap.buckle_type} Buckle \n").bold = True
            blue_index += 1 
            
        for strap in non_urgent_red_straps_NORMLE_BUCKLE:
            p2.add_run(f"{red_index}R | {strap.qty}X '{strap.name}' \n") 
            red_index += 1 

        for strap in non_urgent_red_straps_OLD_BUCKLE:
            p2.add_run(f"{red_index}R | {strap.qty}X '{strap.name}' {strap.buckle_type} Buckle \n")
            red_index += 1
        
        if self.strap_colour is None:
            document_title = 'All.docx' 
        else:
            document_title = f'{self.strap_colour}.docx'

        document.save(document_title) 